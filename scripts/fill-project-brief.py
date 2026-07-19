"""Fill Project Brief template with UMKMan content, preserving Word formatting."""
from docx import Document

src = r"D:\umkm-an-modal-components\Salinan dari Project Brief Template.docx"
out = r"D:\umkm-an-modal-components\UMKMan-Project-Brief.docx"

doc = Document(src)


def set_para_text(paragraph, text: str) -> None:
    """Replace paragraph text while preserving first-run formatting."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r.text = ""


# Problem Statement (P3–P4)
set_para_text(
    doc.paragraphs[3],
    "UMKM di Indonesia (F&B, fashion, jasa, retail, dan sejenisnya) sering kesulitan "
    "mengelola operasional digital sehari-hari: mencatat penjualan dan stok masih manual, "
    "membuat caption & poster konten memakan waktu, serta menjawab chat pelanggan di "
    "WhatsApp berulang. Akibatnya, omzet digital lambat tumbuh dan pemilik usaha kehabisan "
    "waktu untuk pekerjaan berulang.",
)
set_para_text(
    doc.paragraphs[4],
    "Yang terdampak adalah pemilik usaha mikro–kecil multi-klien yang butuh tools sederhana, "
    "multi-tenant (data tiap usaha terisolasi), dan bisa dipakai langsung di web tanpa tim IT. "
    "Masalah ini penting karena digitalisasi UMKM menuntut efisiensi konten, stok, transaksi, "
    "dan layanan pelanggan dalam satu alur terintegrasi.",
)

# Deskripsi Produk (P6–P7)
set_para_text(
    doc.paragraphs[6],
    "UMKMan (UMKM Naik Kelas, Ditenagai Kecerdasan Buatan) adalah platform web multi-tenant "
    "untuk mengelola usaha: dashboard, katalog & stok, pencatatan transaksi, generate caption/"
    "poster AI, asisten AI, laporan keuangan, chatbot WhatsApp (GOWA), serta posting Instagram "
    "(Graph API resmi dan opsi demo private API).",
)
set_para_text(
    doc.paragraphs[7],
    "Produk menyelesaikan masalah dengan menggabungkan data usaha live (PostgreSQL), AI "
    "(Gemini untuk teks, Genity untuk poster), gateway WhatsApp production, dan storage "
    "permanen (Vercel Blob), sehingga setiap client hanya melihat data usahanya sendiri dan "
    "bisa beroperasi end-to-end dari dashboard hingga channel pemasaran.",
)

# Fitur & Teknologi (P9–P13) — ganti contoh template
set_para_text(
    doc.paragraphs[9],
    "Fitur utama (multi-tenant — data mengikuti brand/katalog usaha aktif, tidak hardcode merek demo):",
)
set_para_text(
    doc.paragraphs[10],
    "Multi-tenant dashboard: katalog, restock stok, transaksi AI, konten caption/poster AI, "
    "asisten AI, laporan PDF omzet & modal.",
)
set_para_text(
    doc.paragraphs[11],
    "Channel: WhatsApp chatbot (GOWA + webhook AI), Instagram Graph/assisted/private demo; "
    "auth email + Google OAuth.",
)
set_para_text(
    doc.paragraphs[12],
    "Teknologi: Next.js, React, TypeScript, Tailwind CSS; PostgreSQL + Prisma; Google Gemini "
    "+ GenityBoost; Vercel + Vercel Blob; Railway (Postgres, GOWA, aiograpi-rest).",
)
set_para_text(
    doc.paragraphs[13],
    "Prinsip: fleksibel untuk F&B, fashion, jasa, retail, dan jenis UMKM lain. AI, stok, WA, "
    "dan IG selalu terikat usaha/client aktif.",
)

# Cara Penggunaan (P15–P19)
set_para_text(
    doc.paragraphs[15],
    "Alur pengguna: (1) Buka https://umkman.vercel.app → daftar/login (email atau Google). "
    "(2) Buat usaha (brand, pemilik, kota, kategori) — satu akun bisa banyak usaha/client.",
)
set_para_text(
    doc.paragraphs[16],
    "(3) Pengaturan → Katalog & Stok: tambah produk, restock (+stok / set stok), atur harga. "
    "(4) Catat transaksi (teks AI atau manual) → stok berkurang; restock di Katalog.",
)
set_para_text(
    doc.paragraphs[17],
    "(5) Buat konten: generate caption/poster AI → simpan draft → Post ke Instagram. "
    "(6) Hubungkan WhatsApp (GOWA) untuk chatbot AI; hubungkan Instagram per usaha.",
)
set_para_text(
    doc.paragraphs[18],
    "(7) Asisten AI / unduh laporan PDF untuk insight omzet & modal usaha aktif. "
    "Semua data terisolasi per tenant (client tidak saling melihat data).",
)
set_para_text(
    doc.paragraphs[19],
    "Akses produk: https://umkman.vercel.app — daftar akun baru atau Google OAuth. "
    "Setiap akun mulai kosong (data usaha demo/seed tidak tercampur ke client).",
)

# Informasi Pendukung (P21–P25)
set_para_text(
    doc.paragraphs[21],
    "Studi kasus arah produk: multi-client UMKM (F&B/kopi, retail, jasa, fashion) — bukan "
    "hardcode satu merek; AI & stok mengikuti katalog usaha aktif masing-masing client.",
)
set_para_text(
    doc.paragraphs[22],
    "Dokumentasi teknis di repo: docs/RAILWAY-MCP.md, docs/GOWA.md, docs/VERCEL.md; "
    "stack Next.js + Prisma + Railway (Postgres, GOWA, aiograpi) + Vercel Blob.",
)
set_para_text(
    doc.paragraphs[23],
    "Rencana pengembangan: OAuth Instagram resmi (UX), notifikasi stok menipis, multi-device "
    "WA per cabang, observability log production, dan peningkatan laporan keuangan.",
)
set_para_text(
    doc.paragraphs[24],
    "Tim: Naufal Muhammad Dzaka’ — product owner / full-stack developer (arsitektur multi-tenant, "
    "integrasi AI, WhatsApp, Instagram, deploy Vercel–Railway).",
)
set_para_text(
    doc.paragraphs[25],
    "Production live di Vercel; gateway always-on di Railway. Seed lokal hanya untuk development, "
    "bukan data client production.",
)

doc.save(out)
print("Saved:", out)

# Verify
doc2 = Document(out)
for i, p in enumerate(doc2.paragraphs):
    t = p.text.strip()
    if t:
        print(f"P{i}: {t[:110]}")
print("TABLE:", [c.text for c in doc2.tables[0].rows[1].cells])
